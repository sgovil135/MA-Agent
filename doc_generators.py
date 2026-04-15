"""
Word document generators for Score Memo, Diligence Questions, and LOI Draft.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _init_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


# ── Score Memo ────────────────────────────────────────────────────────────

def build_score_memo(scorecard: dict[str, Any]) -> bytes:
    """Build a Score Memo .docx from the scorecard dict returned by OpenAI."""
    doc = _init_doc()

    company = scorecard.get("company_name", "Unknown Company")
    total = scorecard.get("total_score", 0)
    verdict = scorecard.get("verdict", "")

    doc.add_heading(f"{company} — Deal Scorecard", level=1)

    # Business description
    doc.add_heading("Business Description", level=2)
    doc.add_paragraph(scorecard.get("business_description", "N/A"))

    # Scores
    doc.add_heading("Scores", level=2)
    scores = scorecard.get("scores", {})
    for dim_key, dim_label in [
        ("autonomy", "Autonomy"),
        ("cash_flow_quality", "Cash Flow Quality"),
        ("growth_reality", "Growth Reality"),
        ("downside_risk", "Downside Risk"),
        ("strategic_fit", "Strategic Fit"),
    ]:
        dim = scores.get(dim_key, {})
        score_val = dim.get("score", "?")
        reason = dim.get("reason", "")
        doc.add_paragraph(f"{dim_label}: {score_val}/5 — {reason}", style="List Bullet")

    # Total + Verdict
    doc.add_heading("Total Score & Verdict", level=2)
    doc.add_paragraph(f"Total: {total}/25 — {verdict}")

    # Valuation range
    doc.add_heading("Valuation Range", level=2)
    val = scorecard.get("valuation_range", {})
    low = val.get("low", "?")
    mid = val.get("mid", "?")
    high = val.get("high", "?")
    mult = val.get("multiple_range", "?")
    doc.add_paragraph(f"{low} – {high} ({mult})")
    doc.add_paragraph(f"Mid-case: {mid}")

    # Red flags
    doc.add_heading("Red Flags", level=2)
    red_flags = scorecard.get("red_flags", [])
    if red_flags:
        for flag in red_flags:
            doc.add_paragraph(flag, style="List Bullet")
    else:
        doc.add_paragraph("None identified")

    # Open questions
    doc.add_heading("Open Questions", level=2)
    questions = scorecard.get("open_questions", [])
    for q in questions:
        doc.add_paragraph(q, style="List Number")

    # Executive summary
    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(scorecard.get("executive_summary", "N/A"))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Diligence Questions ──────────────────────────────────────────────────

def build_diligence_doc(company_name: str, questions: dict[str, list[str]]) -> bytes:
    """Build a Diligence Questions .docx from the categorized question dict."""
    doc = _init_doc()

    doc.add_heading(f"{company_name} — Due Diligence Questions", level=1)

    category_labels = {
        "financial": "Financial",
        "legal": "Legal",
        "operational": "Operational",
        "commercial": "Commercial",
        "integration": "Integration",
    }

    for cat_key, cat_label in category_labels.items():
        cat_questions = questions.get(cat_key, [])
        if not cat_questions:
            continue
        doc.add_heading(cat_label, level=2)
        for i, q in enumerate(cat_questions, 1):
            doc.add_paragraph(f"{i}. {q}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── LOI Draft ─────────────────────────────────────────────────────────────

def build_loi_doc(company_name: str, loi_text: str) -> bytes:
    """Build an LOI Draft .docx from the plain text LOI."""
    doc = _init_doc()

    doc.add_heading(f"Letter of Intent — {company_name}", level=1)

    # Split by double newlines into paragraphs, preserving section headers
    sections = loi_text.strip().split("\n\n")
    for section in sections:
        lines = section.strip().split("\n")
        if not lines:
            continue

        first_line = lines[0].strip()

        # Detect section headers (ALL CAPS lines or lines ending with colon)
        if first_line.isupper() or (first_line.endswith(":") and len(first_line) < 60):
            doc.add_heading(first_line.rstrip(":"), level=2)
            body = "\n".join(lines[1:]).strip()
            if body:
                doc.add_paragraph(body)
        else:
            doc.add_paragraph(section.strip())

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
