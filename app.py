import io
import os
import re
from datetime import datetime, timezone

import requests
from docx import Document
from docx.shared import Pt, Inches
from fastapi import FastAPI, Request
from openpyxl import Workbook, load_workbook
from slack_bolt import App
from slack_bolt.adapter.fastapi import SlackRequestHandler
from openai import OpenAI

api = FastAPI()

SLACK_BOT_TOKEN = os.getenv("SLACK_BOT_TOKEN")
SLACK_SIGNING_SECRET = os.getenv("SLACK_SIGNING_SECRET")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

MICROSOFT_CLIENT_ID = os.getenv("MICROSOFT_CLIENT_ID")
MICROSOFT_CLIENT_SECRET = os.getenv("MICROSOFT_CLIENT_SECRET")
MICROSOFT_TENANT_ID = os.getenv("MICROSOFT_TENANT_ID")
SHAREPOINT_SITE_ID = os.getenv("SHAREPOINT_SITE_ID")

openai_client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

# ── Conversation state for follow-up questions ──────────────────────────
# Key: (channel_id, user_id) → deal context dict
pending_deals = {}

# ── Microsoft Graph helpers ─────────────────────────────────────────────

def get_graph_token():
    url = f"https://login.microsoftonline.com/{MICROSOFT_TENANT_ID}/oauth2/v2.0/token"
    resp = requests.post(url, data={
        "client_id": MICROSOFT_CLIENT_ID,
        "client_secret": MICROSOFT_CLIENT_SECRET,
        "scope": "https://graph.microsoft.com/.default",
        "grant_type": "client_credentials",
    }, timeout=30)
    resp.raise_for_status()
    return resp.json()["access_token"]


def graph_headers():
    return {"Authorization": f"Bearer {get_graph_token()}"}


def _drive_url():
    return f"https://graph.microsoft.com/v1.0/sites/{SHAREPOINT_SITE_ID}/drive"


def ensure_folder(folder_path):
    """Create folder if it doesn't exist. folder_path like 'Deal Flow/Pipeline'."""
    parts = folder_path.strip("/").split("/")
    current = ""
    headers = graph_headers()
    for part in parts:
        parent = f"root:/{current}" if current else "root"
        current = f"{current}/{part}" if current else part
        check_url = f"{_drive_url()}/root:/{current}"
        r = requests.get(check_url, headers=headers, timeout=15)
        if r.status_code == 404:
            create_url = f"{_drive_url()}/items/{parent}/children"
            body = {"name": part, "folder": {}, "@microsoft.graph.conflictBehavior": "fail"}
            requests.post(create_url, headers=headers, json=body, timeout=15)


def upload_file(folder_path, file_name, content_bytes, content_type="application/octet-stream"):
    """Upload a file to SharePoint. Returns the item metadata or None."""
    ensure_folder(folder_path)
    headers = graph_headers()
    headers["Content-Type"] = content_type
    safe_name = file_name.replace("/", "-")
    url = f"{_drive_url()}/root:/{folder_path}/{safe_name}:/content"
    r = requests.put(url, headers=headers, data=content_bytes, timeout=60)
    r.raise_for_status()
    return r.json()


def download_file(file_path):
    """Download a file from SharePoint. Returns bytes or None."""
    headers = graph_headers()
    url = f"{_drive_url()}/root:/{file_path}:/content"
    r = requests.get(url, headers=headers, timeout=60)
    if r.status_code == 200:
        return r.content
    return None


def file_exists(file_path):
    """Check if a file exists in SharePoint. Returns metadata dict or None."""
    headers = graph_headers()
    url = f"{_drive_url()}/root:/{file_path}"
    r = requests.get(url, headers=headers, timeout=15)
    if r.status_code == 200:
        return r.json()
    return None


def find_existing_memo(company_name):
    """Search Deal Flow/ for an existing memo with this company name."""
    headers = graph_headers()
    for folder in ["Deal Flow/Pipeline", "Deal Flow/Active", "Deal Flow/Passed"]:
        url = f"{_drive_url()}/root:/{folder}:/children"
        r = requests.get(url, headers=headers, timeout=15)
        if r.status_code != 200:
            continue
        for item in r.json().get("value", []):
            name = item.get("name", "")
            if company_name.lower() in name.lower() and name.endswith(".docx"):
                created = item.get("createdDateTime", "unknown")
                return {"name": name, "folder": folder, "created": created}
    return None


def move_file(old_path, new_folder, file_name):
    """Move a file from old_path to new_folder."""
    headers = graph_headers()
    # Get the item ID
    url = f"{_drive_url()}/root:/{old_path}"
    r = requests.get(url, headers=headers, timeout=15)
    if r.status_code != 200:
        return False
    item_id = r.json()["id"]
    # Get the new parent folder ID
    ensure_folder(new_folder)
    folder_url = f"{_drive_url()}/root:/{new_folder}"
    fr = requests.get(folder_url, headers=headers, timeout=15)
    if fr.status_code != 200:
        return False
    parent_id = fr.json()["id"]
    # Move
    patch_url = f"{_drive_url()}/items/{item_id}"
    body = {"parentReference": {"id": parent_id}, "name": file_name}
    mr = requests.patch(patch_url, headers=headers, json=body, timeout=15)
    return mr.status_code == 200


# ── Word doc generation ─────────────────────────────────────────────────

def build_deal_memo_docx(deal):
    """Build a .docx from a deal dict. Returns bytes."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(f"{deal['company_name']} — Deal Memo", level=1)
    doc.add_paragraph(f"Date: {deal['date']}")

    doc.add_heading("Business Description", level=2)
    doc.add_paragraph(deal.get("business", "N/A"))

    doc.add_heading("Financials", level=2)
    doc.add_paragraph(deal.get("financials", "N/A"))

    doc.add_heading("Scores", level=2)
    for score_line in deal.get("scores", []):
        doc.add_paragraph(score_line, style="List Bullet")

    doc.add_heading("Total Score & Verdict", level=2)
    doc.add_paragraph(deal.get("total_line", "N/A"))

    doc.add_heading("Valuation Range", level=2)
    doc.add_paragraph(deal.get("valuation", "N/A"))

    doc.add_heading("Open Questions", level=2)
    for q in deal.get("open_questions", []):
        doc.add_paragraph(q, style="List Number")

    doc.add_heading("Red Flags", level=2)
    doc.add_paragraph(deal.get("red_flags", "None identified"))

    doc.add_heading("Source", level=2)
    doc.add_paragraph(deal.get("source", "Pending"))

    doc.add_heading("Saagar's Reaction", level=2)
    doc.add_paragraph(deal.get("reaction", "Pending"))

    doc.add_heading("Decision Status", level=2)
    doc.add_paragraph(deal.get("status", "Pipeline"))

    doc.add_heading("Notes", level=2)
    doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Excel master log ────────────────────────────────────────────────────

EXCEL_PATH = "Deal Flow/_Master Deal Log.xlsx"
EXCEL_COLUMNS = [
    "Company", "Date", "Sector", "Source", "Revenue (LTM)", "EBITDA (LTM)",
    "Score", "Verdict", "Decision", "Saagar's Take", "Memo Link",
]


def append_excel_row(row_data):
    """Download the master log, append a row, re-upload."""
    existing = download_file(EXCEL_PATH)
    if existing:
        wb = load_workbook(io.BytesIO(existing))
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "Deal Log"
        ws.append(EXCEL_COLUMNS)

    ws.append([row_data.get(c, "") for c in EXCEL_COLUMNS])

    buf = io.BytesIO()
    wb.save(buf)
    upload_file("Deal Flow", "_Master Deal Log.xlsx", buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def update_excel_row(company_name, column, value):
    """Find a row by company name and update a specific column."""
    existing = download_file(EXCEL_PATH)
    if not existing:
        return
    wb = load_workbook(io.BytesIO(existing))
    ws = wb.active
    col_idx = None
    for i, cell in enumerate(ws[1], 1):
        if cell.value == column:
            col_idx = i
            break
    if not col_idx:
        return
    for row in ws.iter_rows(min_row=2):
        if row[0].value and company_name.lower() in str(row[0].value).lower():
            row[col_idx - 1].value = value
            break
    buf = io.BytesIO()
    wb.save(buf)
    upload_file("Deal Flow", "_Master Deal Log.xlsx", buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ── Scorecard parser ────────────────────────────────────────────────────

def parse_scorecard(text):
    """Extract structured data from the OpenAI scorecard output."""
    deal = {"raw_scorecard": text}

    # Company name from first line
    m = re.search(r"\*(.+?)(?:\s*—\s*Deal Scorecard)\*", text)
    deal["company_name"] = m.group(1).strip() if m else "Unknown Company"

    # Business
    m = re.search(r"\*Business\*\n(.+?)(?=\n\*)", text, re.DOTALL)
    deal["business"] = m.group(1).strip() if m else ""

    # Financials
    m = re.search(r"\*Financials\*\n(.+?)(?=\n\*)", text, re.DOTALL)
    deal["financials"] = m.group(1).strip() if m else ""

    # Individual scores
    scores = []
    for line in text.split("\n"):
        line_stripped = line.strip().lstrip("> ")
        if re.match(r"(Autonomy|Cash Flow Quality|Growth Reality|Downside Risk|Strategic Fit):", line_stripped):
            scores.append(line_stripped)
    deal["scores"] = scores

    # Total line
    m = re.search(r"\*Total:(.+?)\*", text)
    deal["total_line"] = f"Total:{m.group(1)}" if m else ""

    # Extract numeric score
    m = re.search(r"Total:\s*([\d.]+)/25", text)
    deal["score"] = m.group(1) if m else ""

    # Verdict
    verdict = ""
    if "GO FOR IT" in text.upper():
        verdict = "Go for it"
    elif "PROCEED CONSERVATIVELY" in text.upper():
        verdict = "Proceed Conservatively"
    elif "PASS" in text.upper():
        verdict = "Pass"
    deal["verdict"] = verdict

    # Valuation
    m = re.search(r"\*Valuation Range\*\n(.+?)(?=\n\*|\Z)", text, re.DOTALL)
    deal["valuation"] = m.group(1).strip() if m else ""

    # Open questions
    questions = re.findall(r"\d+\.\s+(.+)", text[text.find("Open Questions"):] if "Open Questions" in text else "")
    deal["open_questions"] = questions[:5]

    # Revenue/EBITDA for Excel (last year)
    rev_match = re.search(r"Revenue:.*?([\$\d,.]+[MmBb]?)\s*$", deal["financials"], re.MULTILINE)
    deal["revenue_ltm"] = rev_match.group(1) if rev_match else ""
    ebitda_match = re.search(r"EBITDA:.*?([\$\d,.]+[MmBb]?)\s*$", deal["financials"], re.MULTILINE)
    deal["ebitda_ltm"] = ebitda_match.group(1) if ebitda_match else ""

    # Sector guess from business description
    deal["sector"] = ""
    deal["red_flags"] = ""

    deal["date"] = datetime.now(timezone.utc).strftime("%Y-%m-%d")

    return deal


# ── SharePoint save logic ──────────────────────────────────────────────

def save_deal_to_sharepoint(deal, say, channel):
    """Create Word doc + Excel row. Returns True on success."""
    company = deal["company_name"]
    date = deal["date"]

    try:
        # Check for duplicates
        existing = find_existing_memo(company)
        if existing:
            deal["_duplicate"] = existing
            return "duplicate_found"

        # Build and upload Word doc
        doc_bytes = build_deal_memo_docx(deal)
        doc_name = f"{company} — Deal Memo.docx"
        result = upload_file("Deal Flow/Pipeline", doc_name, doc_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        memo_link = result.get("webUrl", "")
        deal["memo_link"] = memo_link

        # Append Excel row
        append_excel_row({
            "Company": company,
            "Date": date,
            "Sector": deal.get("sector", ""),
            "Source": "",
            "Revenue (LTM)": deal.get("revenue_ltm", ""),
            "EBITDA (LTM)": deal.get("ebitda_ltm", ""),
            "Score": deal.get("score", ""),
            "Verdict": deal.get("verdict", ""),
            "Decision": "Pipeline",
            "Saagar's Take": "",
            "Memo Link": memo_link,
        })

        say(text=f"Saved to SharePoint: `Deal Flow/Pipeline/{doc_name}` and logged in Master Deal Log.", channel=channel)
        return "saved"

    except Exception as e:
        say(text=f"SharePoint save error: {str(e)}\n(Scorecard is still in chat above — nothing lost.)", channel=channel)
        return "error"


def handle_duplicate_response(deal, user_reply, say, channel):
    """Handle user's response to duplicate detection."""
    reply_lower = user_reply.strip().lower()
    company = deal["company_name"]
    date = deal["date"]
    existing = deal.get("_duplicate", {})

    try:
        if "duplicate" in reply_lower:
            say(text="Got it — skipping the save. Scorecard is still above for reference.", channel=channel)
            return None  # Done, no follow-up needed

        elif "updated" in reply_lower or "update" in reply_lower:
            doc_name = f"{company} — Deal Memo (Updated {date}).docx"

            # Compare scorecards if we can get the old doc
            deal_with_changes = dict(deal)
            old_bytes = download_file(f"{existing['folder']}/{existing['name']}")
            if old_bytes:
                comparison = compare_with_openai(old_bytes, deal["raw_scorecard"])
                if comparison:
                    deal_with_changes["what_changed"] = comparison

            doc_bytes = build_deal_memo_docx_with_changes(deal_with_changes)
            result = upload_file("Deal Flow/Pipeline", doc_name, doc_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            memo_link = result.get("webUrl", "")

            append_excel_row({
                "Company": company,
                "Date": date,
                "Sector": deal.get("sector", ""),
                "Source": "",
                "Revenue (LTM)": deal.get("revenue_ltm", ""),
                "EBITDA (LTM)": deal.get("ebitda_ltm", ""),
                "Score": deal.get("score", ""),
                "Verdict": deal.get("verdict", ""),
                "Decision": "Pipeline",
                "Saagar's Take": "",
                "Memo Link": memo_link,
            })
            # Note on old row
            update_excel_row(company, "Saagar's Take", f"Updated CIM received {date}")

            say(text=f"Saved updated memo: `Deal Flow/Pipeline/{doc_name}`", channel=channel)
            return "ask_source"

        elif "new version" in reply_lower or "v2" in reply_lower:
            doc_name = f"{company} — Deal Memo v2.docx"
            doc_bytes = build_deal_memo_docx(deal)
            result = upload_file("Deal Flow/Pipeline", doc_name, doc_bytes,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            memo_link = result.get("webUrl", "")

            append_excel_row({
                "Company": company,
                "Date": date,
                "Sector": deal.get("sector", ""),
                "Source": "",
                "Revenue (LTM)": deal.get("revenue_ltm", ""),
                "EBITDA (LTM)": deal.get("ebitda_ltm", ""),
                "Score": deal.get("score", ""),
                "Verdict": deal.get("verdict", ""),
                "Decision": "Pipeline",
                "Saagar's Take": "",
                "Memo Link": memo_link,
            })

            say(text=f"Saved as new version: `Deal Flow/Pipeline/{doc_name}`", channel=channel)
            return "ask_source"

        else:
            say(text="Didn't catch that — reply *duplicate*, *updated CIM*, or *new version*.", channel=channel)
            return "duplicate_pending"

    except Exception as e:
        say(text=f"SharePoint save error: {str(e)}", channel=channel)
        return "ask_source"


def compare_with_openai(old_docx_bytes, new_scorecard_text):
    """Use OpenAI to compare old and new scorecards."""
    if not openai_client:
        return None
    try:
        uploaded = openai_client.files.create(
            file=("old_memo.docx", io.BytesIO(old_docx_bytes), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            purpose="user_data",
        )
        response = openai_client.responses.create(
            model="gpt-4.1-mini",
            input=[
                {
                    "role": "user",
                    "content": [
                        {"type": "input_file", "file_id": uploaded.id},
                        {"type": "input_text", "text": (
                            f"Here is the NEW scorecard:\n\n{new_scorecard_text}\n\n"
                            "Compare the old deal memo (attached) with this new scorecard. "
                            "Summarize meaningful differences in scores, financials, or risk "
                            "assessment in 3-5 bullet points. Be specific about what changed."
                        )},
                    ],
                }
            ],
        )
        return response.output_text
    except Exception:
        return None


def build_deal_memo_docx_with_changes(deal):
    """Build a .docx that includes a 'What Changed' section at the top."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(f"{deal['company_name']} — Deal Memo (Updated)", level=1)
    doc.add_paragraph(f"Date: {deal['date']}")

    if deal.get("what_changed"):
        doc.add_heading("What Changed", level=2)
        doc.add_paragraph(deal["what_changed"])

    doc.add_heading("Business Description", level=2)
    doc.add_paragraph(deal.get("business", "N/A"))

    doc.add_heading("Financials", level=2)
    doc.add_paragraph(deal.get("financials", "N/A"))

    doc.add_heading("Scores", level=2)
    for score_line in deal.get("scores", []):
        doc.add_paragraph(score_line, style="List Bullet")

    doc.add_heading("Total Score & Verdict", level=2)
    doc.add_paragraph(deal.get("total_line", "N/A"))

    doc.add_heading("Valuation Range", level=2)
    doc.add_paragraph(deal.get("valuation", "N/A"))

    doc.add_heading("Open Questions", level=2)
    for q in deal.get("open_questions", []):
        doc.add_paragraph(q, style="List Number")

    doc.add_heading("Red Flags", level=2)
    doc.add_paragraph(deal.get("red_flags", "None identified"))

    doc.add_heading("Source", level=2)
    doc.add_paragraph(deal.get("source", "Pending"))

    doc.add_heading("Saagar's Reaction", level=2)
    doc.add_paragraph(deal.get("reaction", "Pending"))

    doc.add_heading("Decision Status", level=2)
    doc.add_paragraph(deal.get("status", "Pipeline"))

    doc.add_heading("Notes", level=2)
    doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def update_word_doc_field(deal, field, value):
    """Re-upload the Word doc with an updated field."""
    company = deal["company_name"]
    folder = deal.get("_folder", "Deal Flow/Pipeline")
    doc_name = deal.get("_doc_name", f"{company} — Deal Memo.docx")
    path = f"{folder}/{doc_name}"

    try:
        existing_bytes = download_file(path)
        if not existing_bytes:
            return
        doc = Document(io.BytesIO(existing_bytes))
        # Find the heading and update the paragraph after it
        target_heading = {"source": "Source", "reaction": "Saagar's Reaction"}.get(field)
        if not target_heading:
            return
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == target_heading and i + 1 < len(doc.paragraphs):
                doc.paragraphs[i + 1].text = value
                break
        buf = io.BytesIO()
        doc.save(buf)
        upload_file(folder, doc_name, buf.getvalue(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception:
        pass


# ── CIM prompt ──────────────────────────────────────────────────────────

CIM_PROMPT = (
    "You are an M&A analyst scoring a Confidential Information Memorandum (CIM).\n\n"
    "SCORING RULES:\n"
    "- Scores must be 1, 1.5, 2, 2.5, 3, 3.5, 4, 4.5, or 5 only. No other values.\n\n"
    "SCORE THESE 5 CATEGORIES:\n"
    "1. Autonomy — does real management exist without the seller?\n"
    "2. Cash Flow Quality — are margins durable and healthy?\n"
    "3. Growth Reality — is the growth path clear and practical?\n"
    "4. Downside Risk — how fragile is the revenue base?\n"
    "5. Strategic Fit — does this fit Cemtrex's specific platform? (see acquirer context below)\n\n"
    "ABOUT THE ACQUIRER (for Strategic Fit scoring):\n"
    "- Cemtrex is a public microcap platform company\n"
    "- Core divisions: AIS (industrial contracting/fabrication, roll-up strategy), "
    "Invocon (aerospace & defense engineering), Vicon (security/surveillance), "
    "VDI (simulation, private, not part of Cemtrex)\n"
    "- Acquisition filters: 30%+ gross margins, $1B+ TAM, recurring/repeat revenue, "
    "operational independence (no babysitting), must de-lever cleanly\n"
    "- Strategic priorities: tuck-ins to AIS, aerospace adjacencies via Invocon, "
    "businesses that run without the owner\n"
    "- Hard passes: broken ops, fake margin stories, key-man dependent, "
    "sectors requiring deep new expertise\n"
    "- Long-term direction: space infrastructure (fuel systems, power systems, logistics, life support)\n"
    "Score Strategic Fit against THIS specific platform — not a generic acquirer.\n"
    "Strategic Fit scoring guide:\n"
    "  5 = Direct tuck-in to AIS or Invocon, accretive immediately\n"
    "  3 = Adjacent — could bolt on but needs a thesis\n"
    "  1 = Totally standalone, new sector, would require hands-on management\n\n"
    "VERDICT:\n"
    "- 18+ → \"Go for it\"\n"
    "- 15–17 → \"Proceed Conservatively\"\n"
    "- Under 15 → \"Pass\"\n\n"
    "IMPORTANT: Use Slack formatting, NOT Markdown. Use *text* for bold (single asterisks). "
    "Do NOT use **text** or any Markdown syntax.\n\n"
    "FORMAT THE RESPONSE EXACTLY LIKE THIS:\n\n"
    "*[Company Name] — Deal Scorecard*\n\n"
    "*Business*\n"
    "[2-3 sentence description]\n\n"
    "*Financials*\n"
    "> Revenue: $X / $X / $X\n"
    "> EBITDA: $X / $X / $X\n\n"
    "*Scores*\n"
    "> Autonomy: X/5 — [one sentence]\n"
    "> Cash Flow Quality: X/5 — [one sentence]\n"
    "> Growth Reality: X/5 — [one sentence]\n"
    "> Downside Risk: X/5 — [one sentence]\n"
    "> Strategic Fit: X/5 — [one sentence]\n\n"
    "*Total: X/25 — [VERDICT IN CAPS]*\n\n"
    "*Valuation Range*\n"
    "$XM – $XM ([X–Xx EBITDA])\n\n"
    "*Open Questions*\n"
    "1. [question]\n"
    "2. [question]\n"
    "3. [question]\n\n"
    "Keep it under 400 words. No fluff. If data is missing, say so and flag it as a risk."
)

# ── Slack app ───────────────────────────────────────────────────────────

slack_app = None
handler = None

if SLACK_BOT_TOKEN and SLACK_SIGNING_SECRET:
    slack_app = App(
        token=SLACK_BOT_TOKEN,
        signing_secret=SLACK_SIGNING_SECRET
    )
    handler = SlackRequestHandler(slack_app)

    @slack_app.event("app_mention")
    def handle_mention(body, say):
        raw_text = body["event"].get("text", "")
        text = re.sub(r"<@\w+>", "", raw_text).strip()

        if not openai_client:
            say("OPENAI_API_KEY is missing in Railway.")
            return

        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant in Slack. Reply briefly and clearly."},
                    {"role": "user", "content": text}
                ]
            )
            say(response.choices[0].message.content)
        except Exception as e:
            say(f"OpenAI error: {str(e)}")

    @slack_app.event("file_shared")
    def handle_file_shared(body, client, say):
        event = body["event"]
        file_id = event["file_id"]
        channel = event.get("channel_id")
        user_id = event.get("user_id", "")

        if not openai_client:
            say(text="OPENAI_API_KEY is missing in Railway.", channel=channel)
            return

        try:
            file_info = client.files_info(file=file_id)
            file_obj = file_info["file"]
            file_name = file_obj["name"]
            mimetype = file_obj.get("mimetype", "")
            file_url = file_obj.get("url_private")

            if mimetype != "application/pdf":
                say(text=f"Got your file: *{file_name}*. Right now I only handle PDFs.", channel=channel)
                return

            if not file_url:
                say(text=f"Could not access private URL for *{file_name}*.", channel=channel)
                return

            say(text=f"Reading *{file_name}*...", channel=channel)

            download_resp = requests.get(
                file_url,
                headers={"Authorization": f"Bearer {SLACK_BOT_TOKEN}"},
                timeout=60,
            )
            download_resp.raise_for_status()

            pdf_bytes = download_resp.content

            uploaded_file = openai_client.files.create(
                file=(file_name, io.BytesIO(pdf_bytes), "application/pdf"),
                purpose="user_data",
            )

            response = openai_client.responses.create(
                model="gpt-4.1-mini",
                input=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "input_file", "file_id": uploaded_file.id},
                            {"type": "input_text", "text": CIM_PROMPT},
                        ],
                    }
                ],
            )

            scorecard_text = response.output_text
            say(text=scorecard_text, channel=channel)

            # Parse scorecard and save to SharePoint
            deal = parse_scorecard(scorecard_text)

            has_sharepoint = all([MICROSOFT_CLIENT_ID, MICROSOFT_CLIENT_SECRET,
                                  MICROSOFT_TENANT_ID, SHAREPOINT_SITE_ID])
            if not has_sharepoint:
                say(text="(SharePoint not configured — scorecard not saved.)", channel=channel)
                return

            result = save_deal_to_sharepoint(deal, say, channel)

            if result == "duplicate_found":
                dup = deal["_duplicate"]
                say(text=(
                    f"*Heads up — I already have a deal memo for {deal['company_name']} "
                    f"from {dup['created'][:10]}.*\n"
                    "What is this — a *duplicate*, an *updated CIM*, or a *new version*?"
                ), channel=channel)
                pending_deals[(channel, user_id)] = {
                    "state": "duplicate_pending",
                    "deal": deal,
                }
                return

            if result == "saved":
                say(text=(
                    "*One quick question* — where did this deal come from? "
                    "(broker name, direct, referral, etc.)"
                ), channel=channel)
                pending_deals[(channel, user_id)] = {
                    "state": "ask_source",
                    "deal": deal,
                }

        except Exception as e:
            say(text=f"PDF analysis error: {str(e)}", channel=channel)

    @slack_app.event("message")
    def handle_message(body, say):
        event = body.get("event", {})
        # Ignore bot messages, file shares, and subtypes we don't care about
        if event.get("bot_id") or event.get("subtype"):
            return

        channel = event.get("channel", "")
        user_id = event.get("user", "")
        text = event.get("text", "").strip()
        key = (channel, user_id)

        if not text:
            return

        # Check for folder move commands
        text_lower = text.lower()
        if any(phrase in text_lower for phrase in ["pass", "moving on", "passing"]):
            # Check if there's a recent deal for this user
            if key in pending_deals and pending_deals[key].get("deal"):
                deal = pending_deals[key]["deal"]
                company = deal["company_name"]
                folder = deal.get("_folder", "Deal Flow/Pipeline")
                doc_name = deal.get("_doc_name", f"{company} — Deal Memo.docx")
                old_path = f"{folder}/{doc_name}"
                if move_file(old_path, "Deal Flow/Passed", doc_name):
                    update_excel_row(company, "Decision", "Passed")
                    say(text=f"Moved *{company}* to `Deal Flow/Passed/`.", channel=channel)
                del pending_deals[key]
                return

        if any(phrase in text_lower for phrase in ["active", "exploring"]):
            if key in pending_deals and pending_deals[key].get("deal"):
                deal = pending_deals[key]["deal"]
                company = deal["company_name"]
                folder = deal.get("_folder", "Deal Flow/Pipeline")
                doc_name = deal.get("_doc_name", f"{company} — Deal Memo.docx")
                old_path = f"{folder}/{doc_name}"
                if move_file(old_path, "Deal Flow/Active", doc_name):
                    update_excel_row(company, "Decision", "Active")
                    say(text=f"Moved *{company}* to `Deal Flow/Active/`.", channel=channel)
                return

        # Handle pending conversation states
        if key not in pending_deals:
            return

        state = pending_deals[key]["state"]
        deal = pending_deals[key]["deal"]

        if state == "duplicate_pending":
            next_state = handle_duplicate_response(deal, text, say, channel)
            if next_state == "ask_source":
                say(text=(
                    "*One quick question* — where did this deal come from? "
                    "(broker name, direct, referral, etc.)"
                ), channel=channel)
                pending_deals[key]["state"] = "ask_source"
            elif next_state == "duplicate_pending":
                pass  # Stay in this state, user needs to reply again
            else:
                del pending_deals[key]

        elif state == "ask_source":
            deal["source"] = text
            update_word_doc_field(deal, "source", text)
            update_excel_row(deal["company_name"], "Source", text)
            say(text="*What's your take?* Reply with your reaction and I'll save it to the deal memo.", channel=channel)
            pending_deals[key]["state"] = "ask_reaction"

        elif state == "ask_reaction":
            deal["reaction"] = text
            update_word_doc_field(deal, "reaction", text)
            update_excel_row(deal["company_name"], "Saagar's Take", text)
            say(text=f"Saved. *{deal['company_name']}* is in the pipeline. Say *pass* or *active* anytime to move it.", channel=channel)
            # Keep the deal in state so folder moves still work
            pending_deals[key]["state"] = "done"

        elif state == "done":
            # Already done — only folder move commands apply (handled above)
            pass


@api.get("/")
def root():
    return {
        "ok": True,
        "has_bot_token": bool(SLACK_BOT_TOKEN),
        "has_signing_secret": bool(SLACK_SIGNING_SECRET),
        "has_openai_key": bool(OPENAI_API_KEY),
        "has_sharepoint": all([MICROSOFT_CLIENT_ID, MICROSOFT_CLIENT_SECRET,
                               MICROSOFT_TENANT_ID, SHAREPOINT_SITE_ID]),
    }


@api.post("/slack/events")
async def slack_events(req: Request):
    if handler is None:
        return {"ok": False, "error": "Slack env vars missing"}
    return await handler.handle(req)
