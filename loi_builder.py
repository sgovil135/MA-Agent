"""
LOI Generator — mirrors Saagar's exact LOI structure and language.

Uses the Nelson and PES LOIs as templates. Boilerplate sections are
word-for-word identical. Variable sections are filled by the LLM.
"""

from __future__ import annotations

import io
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


@dataclass
class LOITerms:
    """All variable fields for an LOI."""
    # Date & addressee
    date: str = ""  # e.g. "April 6, 2026"
    seller_names: str = ""  # e.g. "Mark and Deborah Bohler"
    seller_greeting: str = ""  # e.g. "Dear Mr. and Mrs. Bohler,"
    company_name: str = ""
    company_abbreviation: str = ""  # e.g. "PES", "Nelson"
    company_address_lines: list[str] = field(default_factory=list)

    # Opening paragraphs (LLM-generated, 1-2 paragraphs)
    opening_paragraphs: list[str] = field(default_factory=list)

    # Who is the buyer? "Cemtrex" or "AIS" (for industrial tuck-ins)
    buyer_name: str = "Cemtrex"  # "Cemtrex" or "AIS"

    # Deal terms
    total_consideration: str = ""  # e.g. "$25,000,000"
    cash_at_close: str = ""  # e.g. "$20,000,000"
    note_amount: str = ""  # e.g. "$5,000,000"

    # Earnout (optional)
    has_earnout: bool = False
    earnout_text: str = ""  # Full earnout section if applicable

    # NWC
    nwc_floor: str = "[TO BE DETERMINED DURING DILIGENCE]"

    # Property
    property_text: str = ""

    # Employees
    employees_text: str = ""

    # Exclusivity days
    exclusivity_days: str = "sixty (60)"

    # Closing paragraphs (LLM-generated, 1-2 paragraphs)
    closing_paragraphs: list[str] = field(default_factory=list)


def _set_single_spacing(para):
    """Force single spacing, zero before/after on a paragraph."""
    fmt = para.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _add_para(doc: Document, text: str, font_name: str = "Arial",
              font_size: float = 11.0, bold: bool = False,
              left_indent: float | None = None):
    """Add a paragraph with consistent formatting."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if left_indent is not None:
        para.paragraph_format.left_indent = Inches(left_indent)
    _set_single_spacing(para)
    return para


def _add_empty(doc: Document, indent: float | None = None):
    """Add an empty paragraph."""
    para = doc.add_paragraph()
    _set_single_spacing(para)
    if indent is not None:
        para.paragraph_format.left_indent = Inches(indent)


def _add_labeled_section(doc: Document, label: str, text: str,
                          bold_label: bool = False):
    """Add a section like 'Property:\ttext...' with hanging indent.

    Label starts at left margin (0"), body text wraps at 1.5".
    """
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(1.5)
    para.paragraph_format.first_line_indent = Inches(-1.5)
    _set_single_spacing(para)

    # Label run
    label_run = para.add_run(f"{label}:\t")
    label_run.font.name = "Arial"
    label_run.font.size = Pt(11)
    label_run.font.bold = bold_label

    # Content run
    content_run = para.add_run(text)
    content_run.font.name = "Arial"
    content_run.font.size = Pt(11)
    content_run.font.bold = False

    return para


def build_loi(terms: LOITerms) -> bytes:
    """Build an LOI .docx from the terms. Returns bytes."""
    doc = Document()

    # Page setup: 8.5x11, 1" margins
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Set default font and spacing
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # ── Date ──────────────────────────────────────────────────────────
    _add_para(doc, terms.date)
    _add_empty(doc)

    # ── Addressee ─────────────────────────────────────────────────────
    _add_para(doc, terms.seller_names)
    _add_para(doc, terms.company_name)
    for line in terms.company_address_lines:
        _add_para(doc, line)

    _add_empty(doc)
    _add_empty(doc)

    # ── Subject line ──────────────────────────────────────────────────
    abbrev = terms.company_abbreviation or terms.company_name
    subject = f'Subject:\tLetter of Intent to Acquire {terms.company_name} ("{abbrev}", "Company")'
    _add_para(doc, subject, bold=True)
    _add_empty(doc)

    # ── Greeting ──────────────────────────────────────────────────────
    _add_para(doc, terms.seller_greeting)
    _add_empty(doc)

    # ── Opening paragraphs (LLM-generated) ────────────────────────────
    for para_text in terms.opening_paragraphs:
        if para_text and para_text.strip():
            _add_para(doc, para_text)
            _add_empty(doc)

    # ── Offer line ────────────────────────────────────────────────────
    buyer = terms.buyer_name
    offer_line = (
        f'{buyer} (the "Buyer") is willing to offer the following for all the assets of '
        f'{abbrev}, including accounts receivable, inventory, FFE, equipment; '
        f'and operating liabilities including all trade payables, customer deposits, '
        f'and accrued liabilities pertaining to the business, cash free, debt free:'
    )
    _add_para(doc, offer_line)
    _add_empty(doc)

    # ── Deal terms ────────────────────────────────────────────────────
    _add_para(doc, f"Total Consideration:\t\t\t\t\t{terms.total_consideration}", bold=True)
    _add_empty(doc)

    _add_para(doc, "Cash to Seller:")
    _add_para(doc, f"Cash at close:\t\t\t\t\t{terms.cash_at_close}")
    _add_para(doc, "Cash to be paid at Closing by wire transfer to a bank account designated by the Seller.")
    _add_para(doc, f"Note\t\t\t\t\t\t\t{terms.note_amount}")
    _add_para(doc, (
        "Three-year note with payments quarterly with interest at 6%. "
        "The note would be secured against the assets of the Company, "
        "junior to any potential debt from our lender."
    ))

    # ── Earnout (optional) ────────────────────────────────────────────
    if terms.has_earnout and terms.earnout_text:
        _add_empty(doc)
        # Split earnout text into lines and add each
        for line in terms.earnout_text.strip().split("\n"):
            line = line.strip()
            if line:
                _add_para(doc, line)

    _add_empty(doc)

    # ── Net Working Capital ───────────────────────────────────────────
    nwc_text = (
        f"Net Working Capital - The Company will be delivered at the Closing of the "
        f"transaction with a sufficient amount of net working capital to run its business "
        f"in the ordinary course based on historical data, to be confirmed during due "
        f"diligence, but expected to be no less than {terms.nwc_floor}. "
        f"To be calculated as NWC = ( Current Assets excluding Cash ) - "
        f"(Current Liabilities excluding debt)."
    )
    _add_para(doc, nwc_text, bold=True)
    _add_empty(doc)

    # ── Property ──────────────────────────────────────────────────────
    _add_labeled_section(doc, "Property", terms.property_text)
    _add_empty(doc)

    # ── Employees ─────────────────────────────────────────────────────
    _add_labeled_section(doc, "Employees", terms.employees_text)
    _add_empty(doc)

    # Retain all employees
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(1.5)
    _set_single_spacing(para)
    run = para.add_run("Buyer intends to retain all employees of the Company through closing.")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    _add_empty(doc)

    # ── Assets (boilerplate) ──────────────────────────────────────────
    assets_text = (
        "All assets of the business are to be included as part of the transaction "
        "and for clarity this includes but is not limited to all Seller's work in "
        "progress, equipment, tooling, vehicles, inventory, past projects, know-how, "
        "engineering, installation lists, names, branding, certificates, customer "
        "references, financial records, computer programs & records for all past "
        "sales & estimates, & project engineering, all past company files and "
        "drawings(soft or hard) or records, memos, manuals, designs; Trademarks, "
        "copyrights, and all intellectual property. No material assets are to be "
        "sold off without the consent of the Buyer."
    )
    _add_labeled_section(doc, "Assets", assets_text)
    _add_empty(doc)

    # ── Expenses (boilerplate) ────────────────────────────────────────
    expenses_text = (
        "Each of the Prospective Buyer and the Company shall be responsible for "
        "and bear all its own costs and expenses (including any broker's, finders, "
        "counsel and investment banking fees) incurred in connection with the "
        "Transaction, including expenses of the Buyer's or Seller's Representatives "
        "incurred at any time in connection with pursuing or consummating the Transaction."
    )
    _add_labeled_section(doc, "Expenses", expenses_text)
    _add_empty(doc)

    # ── Due Diligence (boilerplate) ───────────────────────────────────
    dd_text = (
        "Buyer shall provide Seller with a list of requested information for their "
        "due diligence investigation no later than 7-days after the signing of the LOI. "
        "Buyer, and its representatives, shall have full access to Company books and "
        "records, files, equipment, and premises at its discretion, with timely prior "
        "written notice as is reasonable and as approved by the Seller, after the signing "
        "of this LOI and throughout, until the Closing Date."
    )
    _add_labeled_section(doc, "Due Diligence", dd_text)
    _add_empty(doc)

    # ── Disclosure (boilerplate) ──────────────────────────────────────
    disclosure_text = (
        "The Seller acknowledges that the Buyer must disclose information to the "
        "public in the regular course of its business and may be required to disclose "
        "information about future agreements executed between the parties under its "
        "obligations to the Securities Exchange Commission and NASDAQ."
    )
    _add_labeled_section(doc, "Disclosure", disclosure_text)
    _add_empty(doc)

    # ── Purchase Agreement (boilerplate) ──────────────────────────────
    pa_text = (
        "The Buyer and its counsel shall be responsible for preparing the initial "
        "draft of the Asset Purchase Agreement. Both parties will work to execute "
        "the Asset Purchase Agreement within forty-five (45) days or sooner if possible."
    )
    _add_labeled_section(doc, "Purchase Agreement", pa_text)
    _add_empty(doc)

    # ── Exclusivity ───────────────────────────────────────────────────
    excl_text = (
        f"For a period of {terms.exclusivity_days} days after this non-binding Letter "
        f"of Intent (the \"LOI\") is fully executed, the Buyer shall have a period of "
        f"exclusivity and the Seller will not negotiate with any other party with respect "
        f"to this transaction unless this LOI is terminated by either Party without reason "
        f"by providing the other 3 days' notice in writing of such intent. If such notice "
        f"is given, neither party shall be liable to the other on account of having "
        f"terminated this LOI."
    )
    _add_labeled_section(doc, "Exclusivity", excl_text)
    _add_empty(doc)

    # ── Confidentiality (boilerplate) ─────────────────────────────────
    conf_text = (
        "Both parties agree that all proprietary and/or confidential information, "
        "whether written and oral, which is disclosed shall be treated with the utmost "
        "confidentiality. Both parties shall maintain confidentiality of this potential "
        "transaction until this transaction is consummated or terminated."
    )
    _add_labeled_section(doc, "Confidentiality", conf_text)
    _add_empty(doc)

    # ── Closing Date (boilerplate) ────────────────────────────────────
    cd_text = (
        "The parties intend to close the Transaction within 75 days of the signing "
        "of this LOI. Buyer and Seller will work proactively to obtain necessary "
        "approvals required to close the Transaction. Closing date can be extended "
        "with mutual consent of both parties in writing."
    )
    _add_labeled_section(doc, "Closing Date", cd_text)
    _add_empty(doc)

    # ── Conduct of Business (boilerplate) ─────────────────────────────
    cob_text = (
        "Until the Closing, the Prospective Seller shall conduct their business only "
        "in the ordinary course, and shall not engage in any extraordinary transactions, "
        "without the Prospective Buyer's prior consent. Seller will also diligently work "
        "towards securing new business and maintaining backlog."
    )
    _add_labeled_section(doc, "Conduct of Business", cob_text)
    _add_empty(doc)

    # ── Non-Binding (boilerplate) ─────────────────────────────────────
    nb_text = (
        "It is understood that this LOI does not constitute or give rise to any legally "
        "binding commitment, other than exclusivity, on the part of any party. Instead, "
        "it merely sets forth their present intentions with respect to the terms proposed, "
        "which terms may or may not become part of a definitive Asset Purchase Agreement, "
        "as a basis for future negotiations."
    )
    _add_labeled_section(doc, "Non-Binding", nb_text)
    _add_empty(doc)

    # ── Governing Law (boilerplate) ───────────────────────────────────
    gl_text = (
        "This LOI will be governed by and construed under the internal laws of the "
        "State of Delaware applicable to a contract made and performed in that state, "
        "without regard to the choice of law or conflict of law principles. Each party "
        "submits to the exclusive personal jurisdiction of the state and federal courts "
        "located in Delaware and agrees that all actions or disputes related to this LOI "
        "shall be brought in such courts."
    )
    _add_labeled_section(doc, "Governing Law", gl_text)
    _add_empty(doc)

    # ── Closing paragraphs (LLM-generated) ────────────────────────────
    for para_text in terms.closing_paragraphs:
        if para_text and para_text.strip():
            _add_para(doc, para_text)
            _add_empty(doc)

    # ── Offer expiry (boilerplate) ────────────────────────────────────
    _add_para(doc, (
        "If the foregoing accurately sets forth our mutual intentions with respect to "
        "the principal terms of the proposed Transaction, please sign below and return "
        "a copy of this letter to the undersigned. This offer will expire after seven days."
    ))
    _add_empty(doc)

    # ── Contact line (boilerplate) ────────────────────────────────────
    _add_para(doc, (
        "I look forward to hearing from you further. You may also contact me directly "
        "at 516-428-1782 or via email: sgovil@cemtrex.com."
    ))
    _add_empty(doc)
    _add_empty(doc)

    # ── Signature block (boilerplate) ─────────────────────────────────
    _add_para(doc, "Very truly yours,")
    _add_empty(doc)
    _add_para(doc, "Saagar Govil")
    _add_para(doc, "Chairman & CEO")
    _add_para(doc, "Cemtrex Inc.")
    _add_empty(doc)
    _add_empty(doc)

    # ── Acceptance block ──────────────────────────────────────────────
    year = terms.date.split(",")[-1].strip() if "," in terms.date else "2026"
    accept = doc.add_paragraph()
    accept.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_single_spacing(accept)
    run = accept.add_run(f"Accepted on this ______ Day of _________, {year}")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = False

    _add_empty(doc)

    company_accept = doc.add_paragraph()
    company_accept.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_single_spacing(company_accept)
    run = company_accept.add_run(f"{terms.company_name}.")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = False

    _add_empty(doc)
    _add_empty(doc)

    name_line = doc.add_paragraph()
    name_line.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_single_spacing(name_line)
    run = name_line.add_run("Name:")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.bold = False

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── LLM prompt for variable sections ─────────────────────────────────────

LOI_VARIABLES_PROMPT = """\
You are filling in the variable sections of a Letter of Intent for Saagar Govil, \
Chairman & CEO of Cemtrex Inc. (Nasdaq: CETX).

All boilerplate language is already written. You ONLY need to provide the \
variable fields listed below.

CIM SCORING RESULTS:
__SCORECARD_JSON__

FINANCIAL DATA:
__FINANCIAL_JSON__

Provide ONLY this JSON (no markdown, no explanation):

{{
  "date": "Month Day, Year (today's date)",
  "seller_names": "Full names of seller(s) as they appear in the CIM",
  "seller_greeting": "Dear Mr./Mrs./Ms. [LastName],",
  "company_name": "Full legal company name",
  "company_abbreviation": "Short name or abbreviation used in CIM",
  "company_address_lines": ["Street", "City, State ZIP"],
  "buyer_name": "Cemtrex or AIS (use AIS if this is an industrial services tuck-in)",
  "opening_paragraphs": [
    "First paragraph: Introduce Cemtrex (Nasdaq: CETX). If buyer is AIS, mention AIS subsidiary and industrial services since 1984. Describe approach: acquire well-run businesses, preserve what makes them successful, invest long-term. NOT private equity, NOT buy-and-flip. Mention specific synergies if relevant.",
    "Second paragraph (optional): If there is a specific strategic angle (like PES engineering capabilities for AIS clients), add a paragraph about it. Otherwise omit this."
  ],
  "total_consideration": "$X,XXX,XXX",
  "cash_at_close": "$X,XXX,XXX (80% of total)",
  "note_amount": "$X,XXX,XXX (20% of total)",
  "has_earnout": false,
  "earnout_text": "Only if has_earnout is true. Include: total earnout amount, duration, revenue targets by year, gross margin gate, payment tiers (below 85%, 85-89.9%, 90-99.9%, 100%+), measurement terms. Otherwise empty string.",
  "nwc_floor": "$X,XXX,XXX or '[TO BE DETERMINED DURING DILIGENCE]'",
  "property_text": "Either: 'Buyer intends to purchase the real estate...' (if owned) with FMV appraisal within 12 months and interim lease, OR 'Buyer would assume or enter into the current lease.' (if leased). Use info from CIM.",
  "employees_text": "At Closing, [Name] shall enter into an Employment Agreement with the Company as [Title], for [duration]. Compensation at fair market value, negotiated in good faith. If multiple owners, list each with their role.",
  "exclusivity_days": "sixty (60) or ninety (90) depending on deal complexity",
  "closing_paragraphs": [
    "Integration paragraph (optional): If buyer is AIS and there are synergies, describe how the company will operate as a business unit within AIS, maintaining brand and autonomy while gaining AIS platform support. Mention earnout alignment if applicable.",
    "Warm close paragraph: 'We believe Cemtrex can offer a nice home for the Company you have spent [X years/decades] building. Our goal is to preserve what has made the Company successful, support the existing team, and invest thoughtfully to enable continued growth over time. With long-term capital, administrative support where helpful, and a commitment to continuity, we believe [Company] will have the opportunity to grow and thrive as part of the Cemtrex[/AIS] portfolio for many years to come.'"
  ]
}}

IMPORTANT:
- Use the mid-case equity value from the financial data for total consideration
- Cash = 80% of total, Note = 20% of total
- Format all dollar amounts with commas: $25,000,000 not $25M
- Only include earnout if the deal has material performance uncertainty
- Opening and closing paragraphs should match Saagar's voice: warm but direct, \
  emphasizes permanence, continuity, team retention, not changing day-to-day ops
- If seller names/address are not in the CIM, use placeholder: "[SELLER NAME]", "[ADDRESS]"
- Return ONLY the JSON object
"""
